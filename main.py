import os
import markdown
from docx import Document
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from bs4 import BeautifulSoup

def convert_markdown_to_docx(markdown_file):
    # Read the Markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()

    # Convert Markdown to HTML
    html_text = markdown.markdown(markdown_text)

    # Create a new DOCX document
    doc = Document()

    # Parse the HTML text to extract Markdown elements
    soup = BeautifulSoup(html_text, 'html.parser')

    # Convert Markdown elements to DOCX elements
    for element in soup.recursiveChildGenerator():
        if hasattr(element, 'name'):
            if element.name == 'h1':
                doc.add_heading(element.text, level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text, level=2)
            elif element.name == 'h3':
                doc.add_heading(element.text, level=3)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    paragraph = doc.add_paragraph()
                    paragraph.add_run('• ').bold = True
                    paragraph.add_run(li.text)
            elif element.name == 'strong' or element.name == 'b':
                doc.add_paragraph().add_run(element.text).bold = True
            elif element.name == 'em' or element.name == 'i':
                doc.add_paragraph().add_run(element.text).italic = True
            elif element.name == 'code':
                doc.add_paragraph().add_run(element.text).style = 'Code'
            elif element.name == 'a':
                link_text = element.text
                link_url = element['href']
                doc.add_paragraph().add_hyperlink(link_text, link_url)
        elif element.name is None:
            doc.add_paragraph(element)

    # Determine the output DOCX file path
    output_file = os.path.splitext(markdown_file)[0] + '.docx'

    # Save the document as a DOCX file
    doc.save(output_file)

    return output_file

# Create a Tkinter root window
root = Tk()
root.withdraw()  # Hide the root window

# Prompt the user to select the input Markdown file
markdown_file = askopenfilename(
    filetypes=[("Markdown Files", "*.md")],
    title="Select the input Markdown file"
)

if markdown_file:
    # Convert Markdown to DOCX and get the output file path
    output_file = convert_markdown_to_docx(markdown_file)

    print(f"Conversion complete. Output file saved as: {output_file}")
else:
    print("No file selected.")
